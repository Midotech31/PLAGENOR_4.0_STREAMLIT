"""
PLAGENOR Document Service
Auto-generates Platform Notes (DOCX), Analysis Reports (DOCX),
Submission PDFs and Invoice PDFs from official EGTP templates.
Matches the official IBTIKAR-DGRSDT form structure from all DOCX files.
"""
import sys, os, io, hashlib
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import config
from core.audit_engine import log_event

PLAGENOR_BLUE = RGBColor(0x1B, 0x4F, 0x72)
ESSBO_GREEN   = RGBColor(0x1A, 0xBC, 0x9C)
BLACK         = RGBColor(0x00, 0x00, 0x00)
GRAY          = RGBColor(0x66, 0x66, 0x66)

DOCS_DIR = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "generated_docs"
)
os.makedirs(DOCS_DIR, exist_ok=True)


# ── Internal helpers ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _header_row(table, *cols, bg="1B4F72"):
    row = table.rows[0]
    for i, col in enumerate(cols):
        if i >= len(row.cells):
            break
        cell      = row.cells[i]
        cell.text = col
        run       = cell.paragraphs[0].runs[0]
        run.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(9)
        _set_cell_bg(cell, bg)


def _bold_para(doc, text, size=11, color=None):
    p        = doc.add_paragraph()
    run      = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def _section_title(doc, number: str, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    run = p.add_run(f"{number}. {title}")
    run.bold = True
    run.font.size      = Pt(11)
    run.font.color.rgb = PLAGENOR_BLUE
    return p


def _divider(doc):
    p = doc.add_paragraph("─" * 80)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.color.rgb = GRAY
        run.font.size      = Pt(7)


def _set_page_margins(doc):
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)


def _add_footer(doc, text: str):
    footer   = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(text)
    run.font.size      = Pt(7)
    run.font.color.rgb = GRAY


def _build_doc_header(doc, left: str, center: str, right: str,
                      center_size: int = 12):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(8)
    tbl.columns[2].width = Cm(4)

    c0 = tbl.cell(0, 0)
    c0.text = left
    c0.paragraphs[0].runs[0].bold = True
    c0.paragraphs[0].runs[0].font.color.rgb = PLAGENOR_BLUE
    _set_cell_bg(c0, "EAF4FB")

    c1 = tbl.cell(0, 1)
    c1.text = center
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    c1.paragraphs[0].runs[0].bold = True
    c1.paragraphs[0].runs[0].font.size      = Pt(center_size)
    c1.paragraphs[0].runs[0].font.color.rgb = PLAGENOR_BLUE

    c2 = tbl.cell(0, 2)
    c2.text = right
    c2.paragraphs[0].runs[0].font.size = Pt(8)
    _set_cell_bg(c2, "EAF4FB")
    return tbl


def _get_checklist(service_code: str) -> list:
    """Returns the official compliance checklist per DOCX service."""
    base = [
        "Échantillons reçus en bon état (aspect, température)",
        "Formulaire rempli intégralement",
    ]
    checklists = {
        "EGTP-IMT": base + [
            "Quantité minimale d'échantillon respectée",
            "Mode de conservation/transport respecté",
            "Déclaration de risque biologique fournie (si pathogène)",
        ],
        "EGTP-GDE": base + [
            "Quantité minimale d'échantillon respectée",
            "Mode de conservation/transport respecté",
            "Contrôle qualité d'ADN fourni",
            "Rapport de sérologie fourni (si échantillon biologique)",
        ],
        "EGTP-Seq01": base + [
            "Quantité minimale d'échantillon respectée",
        ],
        "EGTP-Seq02": base + [
            "Quantité minimale d'échantillon respectée",
            "Mode de conservation/transport respecté",
            "Contrôle qualité d'ADN fourni",
        ],
        "EGTP-SeqS": base + [
            "Amplicons PCR purs fournis",
            "Résultats QC des amplicons fournis",
            "Amorces de séquençage fournies",
        ],
        "EGTP-PCR": base + [
            "Volume d'ADN matriciel ≥10 µL fourni",
            "Amorces F+R (5 µL, 10 µM) fournies",
            "Séquences des amorces fournies en format clair (5'→3')",
        ],
        "EGTP-CAN": base + [
            "Volume d'échantillon suffisant (≥10 µL)",
        ],
        "EGTP-PS": base + [
            "Séquences des amorces fournies en format clair (5'→3')",
            "Longueur des amorces compatible avec la synthèse standard",
        ],
        "EGTP-Lyoph": base + [
            "Mode de conservation/transport respecté",
            "Récipients adaptés à l'équipement Beta 2-8 LSCplus",
            "Déclaration biosécurité fournie (si pathogène)",
            "Absence de solvants interdits confirmée",
        ],
        "EGTP-Illumina-Microbial-WGS": base + [
            "4 × 5 mL tubes milieu liquide stérile fournis",
            "Mode de conservation/transport respecté",
            "Déclaration de risque biologique fournie (si pathogène)",
        ],
    }
    return checklists.get(service_code, base)


# ── FUNCTION 1: generate_platform_note ───────────────────────────────────────

def generate_platform_note(request: dict, user: dict) -> str:
    """
    Generates an official Platform Note (Note de Plateforme) DOCX.
    Matches IBTIKAR-DGRSDT administrative format.
    Locked after generation — version tracked.
    Returns file path.
    """
    doc      = Document()
    _set_page_margins(doc)

    form_data = request.get("form_data", {})
    svc_code  = request.get("service_code", "N/A")

    # Header
    _build_doc_header(
        doc,
        left   = "IBTIKAR / PLAGENOR",
        center = "NOTE DE PLATEFORME — PLAGENOR-ESSBO",
        right  = f"{svc_code}\nV01\nIBTIKAR",
    )
    doc.add_paragraph()

    # Reference
    ref = doc.add_table(rows=1, cols=2)
    ref.style = "Table Grid"
    ref.cell(0, 0).text = "N° de la demande d'analyse"
    ref.cell(0, 1).text = (
        f"{request['id'][:8].upper()} · "
        f"{datetime.utcnow().strftime('%d/%m/%Y')}"
    )
    _set_cell_bg(ref.cell(0, 0), "D6EAF8")
    ref.cell(0, 0).paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # Section 1 — Requester
    _section_title(doc, "1", "Informations du demandeur")
    req_info  = form_data.get("requester", {})
    req_table = doc.add_table(rows=6, cols=2)
    req_table.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Nom et prénom",       req_info.get("full_name", "")),
        ("Université / École",  req_info.get("institution", "")),
        ("Laboratoire",         req_info.get("laboratory", "")),
        ("Fonction / Poste",    req_info.get("function", "")),
        ("Adresse e-mail",      req_info.get("email", "")),
        ("Numéro de téléphone", req_info.get("phone", "")),
    ]):
        req_table.rows[i].cells[0].text = label
        req_table.rows[i].cells[1].text = str(val)
        req_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(req_table.rows[i].cells[0], "EBF5FB")
    doc.add_paragraph()

    # Section 2 — Analysis info
    _section_title(doc, "2", "Informations relatives à la demande d'analyse")
    analysis_info = form_data.get("analysis_info", {})
    ana_table     = doc.add_table(rows=3, cols=2)
    ana_table.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Cadre de l'analyse",     analysis_info.get("analysis_frame", "")),
        ("Titre du projet",        analysis_info.get("project_title", "")),
        ("Directeur de recherche", analysis_info.get("director", "")),
    ]):
        ana_table.rows[i].cells[0].text = label
        ana_table.rows[i].cells[1].text = str(val)
        ana_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(ana_table.rows[i].cells[0], "EBF5FB")
    doc.add_paragraph()

    # Section 3 — Samples
    _section_title(doc, "3", "Informations sur les échantillons")
    samples = (
        form_data.get("samples") or
        form_data.get("reactions") or
        form_data.get("primer_sets") or []
    )
    if samples:
        col_keys    = [k for k in samples[0].keys() if k != "remarks"][:6]
        col_headers = [k.replace("_", " ").title() for k in col_keys]
        s_table     = doc.add_table(
            rows=len(samples) + 1, cols=len(col_headers) + 2)
        s_table.style = "Table Grid"
        _header_row(s_table, "N°", *col_headers, "Remarques")
        for idx, sample in enumerate(samples):
            row = s_table.rows[idx + 1]
            row.cells[0].text = str(idx + 1)
            for j, key in enumerate(col_keys):
                row.cells[j + 1].text = str(sample.get(key, ""))
            row.cells[-1].text = str(sample.get("remarks", ""))
    else:
        doc.add_paragraph("Aucun échantillon enregistré.")
    doc.add_paragraph()

    # Section 4 — Pricing
    _section_title(doc, "4", "Tarification")
    pricing     = form_data.get("pricing", {})
    price_table = doc.add_table(rows=1, cols=2)
    price_table.style = "Table Grid"
    _header_row(price_table, "Prestation", "Montant (DA)")
    for item in pricing.get("breakdown", []):
        if item.get("amount", 0) > 0:
            r = price_table.add_row()
            r.cells[0].text = item["label"]
            r.cells[1].text = f"{item['amount']:,.0f} DA"
    total_row = price_table.add_row()
    total_row.cells[0].text = "TOTAL ESTIMÉ"
    total_row.cells[1].text = f"{pricing.get('total_dzd', 0):,.0f} DA"
    total_row.cells[0].paragraphs[0].runs[0].bold = True
    total_row.cells[1].paragraphs[0].runs[0].bold = True
    _set_cell_bg(total_row.cells[0], "D5F5E3")
    _set_cell_bg(total_row.cells[1], "D5F5E3")
    doc.add_paragraph()

    # Section 5 — Admin validation block
    _section_title(doc, "5", "Validation de la demande — Cadre réservé PLAGENOR")
    val_table = doc.add_table(rows=6, cols=2)
    val_table.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Opérateur — Nom et prénom", ""),
        ("Date de réception",         ""),
        ("Prix validé (DA)",          ""),
        ("Notes internes",            ""),
        ("Signature",                 ""),
        ("Visa Chef Service Commun",  ""),
    ]):
        val_table.rows[i].cells[0].text = label
        val_table.rows[i].cells[1].text = val
        val_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(val_table.rows[i].cells[0], "FDFEFE")

    # Checklist
    doc.add_paragraph()
    _bold_para(doc, "Checklist de conformité — À remplir par PLAGENOR",
               size=10, color=PLAGENOR_BLUE)
    for item in _get_checklist(svc_code):
        p   = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"☐  {item}")
        run.font.size = Pt(9)

    # Section 6 — Ethical declaration
    doc.add_paragraph()
    _section_title(doc, "6", "Déclaration de responsabilité éthique")
    p   = doc.add_paragraph()
    run = p.add_run(
        "La signature du présent formulaire engage le demandeur à certifier que les "
        "échantillons soumis ont été collectés, manipulés et transférés dans le respect "
        "strict des normes éthiques et réglementaires en vigueur. Il reconnaît être "
        "pleinement responsable de la nature, de l'origine et de l'utilisation de ces "
        "échantillons, y compris de toute implication éthique ou juridique liée à leur "
        "traitement ou à leur analyse."
    )
    run.font.size   = Pt(9)
    run.font.italic = True

    # Signature lines
    doc.add_paragraph()
    sig = doc.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    sig.cell(0, 0).text = "Signature du demandeur"
    sig.cell(0, 1).text = "Date"
    sig.cell(1, 0).text = " " * 40
    sig.cell(1, 1).text = " " * 20
    for cell in sig.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "EBF5FB")

    # Visa des responsables
    doc.add_paragraph()
    _bold_para(doc, "Visa des responsables", size=10, color=PLAGENOR_BLUE)
    visa = doc.add_table(rows=2, cols=2)
    visa.style = "Table Grid"
    visa.cell(0, 0).text = "Visa du Chef du Service Commun"
    visa.cell(0, 1).text = "Visa du Directeur de l'ESSBO"
    visa.cell(1, 0).text = " " * 30
    visa.cell(1, 1).text = " " * 30
    for cell in visa.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "EBF5FB")

    # Footer
    _add_footer(doc,
        f"PLAGENOR — ESSBO Oran | IBTIKAR-DGRSDT | "
        f"Prof. Mohamed Merzoug | "
        f"Generated: {datetime.utcnow().strftime('%d/%m/%Y %H:%M')} UTC"
    )

    # Save
    filename = (
        f"NOTE_PLATEFORME_{svc_code}_{request['id'][:8].upper()}_"
        f"{datetime.utcnow().strftime('%Y%m%d_%H%M')}.docx"
    )
    filepath = os.path.join(DOCS_DIR, filename)
    doc.save(filepath)
    log_event("DOCUMENT", request["id"], "PLATFORM_NOTE_GENERATED",
              user["id"], {"filename": filename, "path": filepath})
    return filepath


# ── FUNCTION 2: generate_report_docx ─────────────────────────────────────────

def generate_report_docx(request: dict, user: dict) -> str:
    """
    Generates the Analysis Report DOCX after analysis completion.
    Called at ANALYSIS_FINISHED state by admin_dashboard.
    Returns file path.
    """
    doc     = Document()
    _set_page_margins(doc)

    form_data = request.get("form_data", {})
    svc_code  = request.get("service_code", "N/A")
    channel   = request.get("channel", "IBTIKAR")
    req_info  = form_data.get("requester", {})
    samples   = (
        form_data.get("samples") or
        form_data.get("reactions") or []
    )

    # Header
    _build_doc_header(
        doc,
        left   = f"PLAGENOR — {channel}",
        center = "RAPPORT D'ANALYSE",
        right  = f"{svc_code}\nIBTIKAR-DGRSDT",
        center_size = 13,
    )
    doc.add_paragraph()

    # Reference
    ref = doc.add_table(rows=1, cols=2)
    ref.style = "Table Grid"
    ref.cell(0, 0).text = "Référence"
    ref.cell(0, 1).text = (
        f"{request['id'][:8].upper()} · "
        f"{datetime.utcnow().strftime('%d/%m/%Y')}"
    )
    _set_cell_bg(ref.cell(0, 0), "D6EAF8")
    ref.cell(0, 0).paragraphs[0].runs[0].bold = True
    doc.add_paragraph()

    # Section 1 — Requester
    _section_title(doc, "1", "Informations du demandeur")
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Nom et prénom",      req_info.get("full_name", "")),
        ("Université / École", req_info.get("institution", "")),
        ("Laboratoire",        req_info.get("laboratory", "")),
        ("Email",              req_info.get("email", "")),
    ]):
        t.rows[i].cells[0].text = label
        t.rows[i].cells[1].text = str(val)
        t.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(t.rows[i].cells[0], "EBF5FB")
    doc.add_paragraph()

    # Section 2 — Analysis summary
    _section_title(doc, "2", "Résumé de l'analyse")
    analysis = form_data.get("analysis_info", {})
    t2 = doc.add_table(rows=3, cols=2)
    t2.style = "Table Grid"
    for i, (label, val) in enumerate([
        ("Service",         svc_code),
        ("Cadre",           analysis.get("analysis_frame", "")),
        ("Titre du projet", analysis.get("project_title", "")),
    ]):
        t2.rows[i].cells[0].text = label
        t2.rows[i].cells[1].text = str(val)
        t2.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(t2.rows[i].cells[0], "EBF5FB")
    doc.add_paragraph()

    # Section 3 — Samples processed
    _section_title(doc, "3", "Échantillons traités")
    if samples:
        col_keys = [k for k in samples[0].keys() if k != "remarks"][:5]
        s_tbl = doc.add_table(rows=len(samples) + 1,
                               cols=len(col_keys) + 2)
        s_tbl.style = "Table Grid"
        _header_row(s_tbl, "N°",
                    *[k.replace("_", " ").title() for k in col_keys],
                    "Remarques")
        for idx, s in enumerate(samples):
            row = s_tbl.rows[idx + 1]
            row.cells[0].text = str(idx + 1)
            for j, key in enumerate(col_keys):
                row.cells[j + 1].text = str(s.get(key, ""))
            row.cells[-1].text = str(s.get("remarks", ""))
    else:
        doc.add_paragraph("Aucun échantillon enregistré.")
    doc.add_paragraph()

    # Section 4 — Results (placeholder for analyst)
    _section_title(doc, "4", "Résultats de l'analyse")
    p   = doc.add_paragraph()
    run = p.add_run(
        "[Les résultats détaillés de l'analyse seront insérés par "
        "l'analyste assigné avant transmission au demandeur.]"
    )
    run.font.italic    = True
    run.font.color.rgb = GRAY
    doc.add_paragraph()

    # Section 5 — Conclusion (placeholder)
    _section_title(doc, "5", "Conclusion et recommandations")
    doc.add_paragraph(
        "[Conclusion et recommandations à compléter par l'analyste.]"
    )
    doc.add_paragraph()

    # Section 6 — Validation
    _section_title(doc, "6", "Validation du rapport")
    val = doc.add_table(rows=4, cols=2)
    val.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("Analyste responsable",  ""),
        ("Date de finalisation",  ""),
        ("Validé par (admin)",    ""),
        ("Date de validation",    ""),
    ]):
        val.rows[i].cells[0].text = label
        val.rows[i].cells[1].text = value
        val.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        _set_cell_bg(val.rows[i].cells[0], "EBF5FB")
    doc.add_paragraph()

    # Signature lines
    sig = doc.add_table(rows=2, cols=2)
    sig.style = "Table Grid"
    sig.cell(0, 0).text = "Visa de l'analyste"
    sig.cell(0, 1).text = "Visa du responsable PLAGENOR"
    sig.cell(1, 0).text = " " * 30
    sig.cell(1, 1).text = " " * 30
    for cell in sig.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        _set_cell_bg(cell, "EBF5FB")

    # Footer
    _add_footer(doc,
        f"PLAGENOR — ESSBO Oran | {channel} | {svc_code} | "
        f"{request['id'][:8].upper()} | "
        f"Generated: {datetime.utcnow().strftime('%d/%m/%Y %H:%M')} UTC"
    )

    # Save
    filename = (
        f"RAPPORT_{svc_code}_{request['id'][:8].upper()}_"
        f"{datetime.utcnow().strftime('%Y%m%d_%H%M')}.docx"
    )
    filepath = os.path.join(DOCS_DIR, filename)
    doc.save(filepath)
    log_event("DOCUMENT", request["id"], "REPORT_DOCX_GENERATED",
              user["id"], {"filename": filename})
    return filepath


# ── FUNCTION 3: generate_submission_pdf ──────────────────────────────────────

def generate_submission_pdf(request: dict) -> bytes:
    """
    Generates a submission confirmation PDF for requester download.
    Returns PDF bytes. Falls back to DOCX bytes if reportlab unavailable.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        )

        buffer    = io.BytesIO()
        doc       = SimpleDocTemplate(
            buffer, pagesize=A4,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
            topMargin=2*cm,    bottomMargin=2*cm
        )
        styles    = getSampleStyleSheet()
        BLUE      = colors.HexColor("#1B4F72")
        LBLUE     = colors.HexColor("#EAF4FB")
        GREEN     = colors.HexColor("#27AE60")
        form_data = request.get("form_data", {})
        pricing   = form_data.get("pricing", {})
        req_info  = form_data.get("requester", {})
        svc_code  = request.get("service_code", "N/A")
        story     = []

        # Title
        story.append(Paragraph(
            "PLAGENOR — Confirmation de Soumission",
            ParagraphStyle("Title", parent=styles["Title"],
                           textColor=BLUE, fontSize=16, spaceAfter=6)))
        story.append(Paragraph(
            f"IBTIKAR-DGRSDT · ESSBO Oran · {svc_code}",
            ParagraphStyle("Sub", parent=styles["Normal"],
                           textColor=colors.grey, fontSize=9, spaceAfter=12)))

        # Reference box
        ref_table = Table([
            ["Référence demande",  request["id"][:8].upper()],
            ["Service",            svc_code],
            ["Date de soumission", datetime.utcnow().strftime("%d/%m/%Y %H:%M UTC")],
            ["Statut",             "SUBMITTED — En attente de validation"],
        ], colWidths=[5*cm, 12*cm])
        ref_table.setStyle(TableStyle([
            ("BACKGROUND",     (0, 0), (0, -1), LBLUE),
            ("FONTNAME",       (0, 0), (0, -1), "Helvetica-Bold"),
            ("TEXTCOLOR",      (0, 0), (0, -1), BLUE),
            ("GRID",           (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("FONTSIZE",       (0, 0), (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 0), (-1, -1),
             [colors.white, colors.HexColor("#F8F9FA")]),
        ]))
        story.append(ref_table)
        story.append(Spacer(1, 0.4*cm))

        # Requester info
        story.append(Paragraph(
            "Informations du demandeur",
            ParagraphStyle("H2", parent=styles["Heading2"],
                           textColor=BLUE, fontSize=11)))
        req_table = Table([
            ["Nom et prénom",       req_info.get("full_name", "")],
            ["Université / École",  req_info.get("institution", "")],
            ["Laboratoire",         req_info.get("laboratory", "")],
            ["Fonction",            req_info.get("function", "")],
            ["Email",               req_info.get("email", "")],
            ["Téléphone",           req_info.get("phone", "")],
        ], colWidths=[5*cm, 12*cm])
        req_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (0, -1), LBLUE),
            ("FONTNAME",   (0, 0), (0, -1), "Helvetica-Bold"),
            ("GRID",       (0, 0), (-1, -1), 0.5, colors.lightgrey),
            ("FONTSIZE",   (0, 0), (-1, -1), 9),
        ]))
        story.append(req_table)
        story.append(Spacer(1, 0.4*cm))

        # Pricing
        story.append(Paragraph(
            "Tarification estimée",
            ParagraphStyle("H2", parent=styles["Heading2"],
                           textColor=BLUE, fontSize=11)))
        price_rows = [["Prestation", "Montant (DA)"]]
        for item in pricing.get("breakdown", []):
            if item.get("amount", 0) > 0:
                price_rows.append([item["label"],
                                   f"{item['amount']:,.0f} DA"])
        price_rows.append(["TOTAL", f"{pricing.get('total_dzd', 0):,.0f} DA"])

        p_table = Table(price_rows, colWidths=[13*cm, 4*cm])
        p_table.setStyle(TableStyle([
            ("BACKGROUND",     (0, 0),  (-1, 0),  BLUE),
            ("TEXTCOLOR",      (0, 0),  (-1, 0),  colors.white),
            ("FONTNAME",       (0, 0),  (-1, 0),  "Helvetica-Bold"),
            ("FONTNAME",       (0, -1), (-1, -1), "Helvetica-Bold"),
            ("BACKGROUND",     (0, -1), (-1, -1), colors.HexColor("#D5F5E3")),
            ("TEXTCOLOR",      (0, -1), (-1, -1), GREEN),
            ("GRID",           (0, 0),  (-1, -1), 0.5, colors.lightgrey),
            ("FONTSIZE",       (0, 0),  (-1, -1), 9),
            ("ROWBACKGROUNDS", (0, 1),  (-1, -2),
             [colors.white, colors.HexColor("#F8F9FA")]),
        ]))
        story.append(p_table)
        story.append(Spacer(1, 0.6*cm))

        # Disclaimer note
        story.append(Paragraph(
            "Le tarif indiqué est estimatif et sujet à validation par l'administrateur "
            "PLAGENOR. Les prix sont susceptibles d'être ajustés en fonction de la "
            "disponibilité des réactifs. Toute modification sera communiquée avant confirmation.",
            ParagraphStyle("Note", parent=styles["Normal"],
                           textColor=colors.grey, fontSize=8)))

        # Footer
        story.append(Spacer(1, 1*cm))
        story.append(Paragraph(
            f"PLAGENOR — ESSBO Oran | IBTIKAR-DGRSDT | "
            f"Prof. Mohamed Merzoug | {datetime.utcnow().strftime('%d/%m/%Y')}",
            ParagraphStyle("Footer", parent=styles["Normal"],
                           textColor=colors.grey, fontSize=7, alignment=1)))

        doc.build(story)
        return buffer.getvalue()

    except ImportError:
        docx_path = generate_platform_note(
            request, {"id": "system", "role": "SYSTEM"})
        with open(docx_path, "rb") as f:
            return f.read()


# ── FUNCTION 4: generate_invoice_pdf ─────────────────────────────────────────

def generate_invoice_pdf(request: dict, invoice: dict, user: dict) -> bytes:
    """
    Generates a GENOCLAB commercial invoice PDF.
    Immutable once generated — tamper detection via SHA-256 hash.
    Returns PDF bytes.
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, HRFlowable
        )

        buffer   = io.BytesIO()
        doc_pdf  = SimpleDocTemplate(
            buffer, pagesize=A4,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
            topMargin=2*cm,    bottomMargin=2.5*cm
        )
        styles       = getSampleStyleSheet()
        BLUE         = colors.HexColor("#1B4F72")
        TEAL         = colors.HexColor("#1ABC9C")
        LBLUE        = colors.HexColor("#EAF4FB")
        GREEN        = colors.HexColor("#27AE60")
        LGRAY        = colors.HexColor("#F8F9FA")

        form_data    = request.get("form_data", {})
        req_info     = form_data.get("requester", {})
        pricing      = form_data.get("pricing", {})
        invoice_num  = invoice.get("invoice_number", "GENOCLAB-INV-0000")
        invoice_date = invoice.get("created_at",
                                   datetime.utcnow().isoformat())[:10]
        svc_code     = request.get("service_code", "N/A")
        vat_rate     = float(invoice.get("vat_rate", 0.19))
        subtotal     = float(invoice.get("subtotal",
                                         pricing.get("total_dzd", 0)))
        vat_amount   = round(subtotal * vat_rate, 2)
        total_ttc    = round(subtotal + vat_amount, 2)
        story        = []

        # Header
        header_data = [[
            Paragraph(
                "<font size='16'><b>PLAGENOR</b></font><br/>"
                "<font size='9' color='#1ABC9C'>GENOCLAB — Division Commerciale</font><br/>"
                "<font size='8' color='grey'>ESSBO · Oran, Algeria</font>",
                styles["Normal"]),
            Paragraph(
                f"<font size='18' color='#1B4F72'><b>FACTURE</b></font><br/>"
                f"<font size='9' color='grey'>N° {invoice_num}</font><br/>"
                f"<font size='9' color='grey'>Date: {invoice_date}</font>",
                ParagraphStyle("Right", parent=styles["Normal"], alignment=2)),
        ]]
        h_table = Table(header_data, colWidths=[9*cm, 8*cm])
        h_table.setStyle(TableStyle([
            ("VALIGN",       (0, 0), (-1, -1), "TOP"),
            ("BACKGROUND",   (0, 0), (-1, -1), LBLUE),
            ("TOPPADDING",   (0, 0), (-1, -1), 12),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 12),
            ("LEFTPADDING",  (0, 0), (-1, -1), 14),
            ("RIGHTPADDING", (0, 0), (-1, -1), 14),
        ]))
        story.append(h_table)
        story.append(Spacer(1, 0.5*cm))

        # Client info
        story.append(Paragraph(
            "<b>FACTURÉ À</b>",
            ParagraphStyle("Label", parent=styles["Normal"],
                           textColor=TEAL, fontSize=8,
                           fontName="Helvetica-Bold", spaceAfter=4)))
        c_table = Table([
            [req_info.get("full_name", ""),
             req_info.get("institution", "")],
            [req_info.get("laboratory", ""),
             req_info.get("email", "")],
        ], colWidths=[8.5*cm, 8.5*cm])
        c_table.setStyle(TableStyle([
            ("FONTSIZE",   (0, 0), (-1, -1), 9),
            ("BACKGROUND", (0, 0), (-1, -1), LGRAY),
            ("GRID",       (0, 0), (-1, -1), 0.3, colors.lightgrey),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(c_table)
        story.append(Spacer(1, 0.4*cm))

        # Service reference line
        story.append(Paragraph(
            f"<b>Référence:</b> {request['id'][:8].upper()} &nbsp;·&nbsp; "
            f"<b>Service:</b> {svc_code} &nbsp;·&nbsp; "
            f"<b>Canal:</b> GENOCLAB",
            ParagraphStyle("Ref", parent=styles["Normal"],
                           fontSize=8, textColor=colors.grey, spaceAfter=8)))
        story.append(HRFlowable(
            width="100%", thickness=1, color=TEAL, spaceAfter=8))

        # Line items
        story.append(Paragraph(
            "<b>DÉTAIL DES PRESTATIONS</b>",
            ParagraphStyle("Label2", parent=styles["Normal"],
                           textColor=BLUE, fontSize=9,
                           fontName="Helvetica-Bold", spaceAfter=6)))

        item_rows = [["Description", "Montant (DA)"]]
        for item in pricing.get("breakdown", []):
            if item.get("amount", 0) > 0:
                item_rows.append([item["label"],
                                  f"{item['amount']:,.2f}"])
        item_rows.append(["", ""])
        item_rows.append(["Sous-total HT",
                          f"{subtotal:,.2f} DA"])
        item_rows.append([f"TVA ({vat_rate*100:.0f}%)",
                          f"{vat_amount:,.2f} DA"])
        item_rows.append(["TOTAL TTC",
                          f"{total_ttc:,.2f} DA"])

        i_table = Table(item_rows, colWidths=[13*cm, 4*cm])
        i_table.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0),  (-1, 0),  BLUE),
            ("TEXTCOLOR",     (0, 0),  (-1, 0),  colors.white),
            ("FONTNAME",      (0, 0),  (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0),  (-1, -1), 9),
            ("GRID",          (0, 0),  (-1, -1), 0.3, colors.lightgrey),
            ("ROWBACKGROUNDS",(0, 1),  (-1, -5),
             [colors.white, LGRAY]),
            ("FONTNAME",      (0, -3), (-1, -3), "Helvetica-Bold"),
            ("LINEABOVE",     (0, -3), (-1, -3), 1, TEAL),
            ("BACKGROUND",    (0, -1), (-1, -1),
             colors.HexColor("#D5F5E3")),
            ("TEXTCOLOR",     (0, -1), (-1, -1), GREEN),
            ("FONTNAME",      (0, -1), (-1, -1), "Helvetica-Bold"),
            ("FONTSIZE",      (0, -1), (-1, -1), 11),
            ("LINEABOVE",     (0, -1), (-1, -1), 1.5, GREEN),
            ("TOPPADDING",    (0, 0),  (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0),  (-1, -1), 5),
        ]))
        story.append(i_table)
        story.append(Spacer(1, 0.5*cm))

        # Payment conditions
        story.append(HRFlowable(
            width="100%", thickness=0.5,
            color=colors.lightgrey, spaceAfter=6))
        story.append(Paragraph(
            "<b>Conditions de paiement:</b> Virement bancaire sous 30 jours. "
            "Tout retard entraîne des pénalités conformément à la réglementation.",
            ParagraphStyle("Pay", parent=styles["Normal"],
                           fontSize=8, textColor=colors.grey, spaceAfter=4)))
        story.append(Paragraph(
            "<b>Coordonnées bancaires:</b> À communiquer sur demande — "
            "contact: mohamed.merzoug.essbo@gmail.com",
            ParagraphStyle("Bank", parent=styles["Normal"],
                           fontSize=8, textColor=colors.grey)))
        story.append(Spacer(1, 0.8*cm))

        # Tamper-detection hash
        hash_input = (
            f"{invoice_num}{invoice_date}{subtotal}"
            f"{vat_amount}{total_ttc}{request['id']}"
        )
        doc_hash = hashlib.sha256(
            hash_input.encode()).hexdigest()[:16].upper()
        story.append(Paragraph(
            f"<font size='7' color='grey'>"
            f"Intégrité: #{doc_hash} — GENOCLAB · PLAGENOR-ESSBO · "
            f"{invoice_num} · "
            f"{datetime.utcnow().strftime('%d/%m/%Y %H:%M')} UTC — "
            f"Document généré automatiquement et faisant foi."
            f"</font>",
            ParagraphStyle("Hash", parent=styles["Normal"],
                           fontSize=7, textColor=colors.grey, alignment=1)))

        doc_pdf.build(story)
        pdf_bytes = buffer.getvalue()

        log_event("DOCUMENT", request["id"], "INVOICE_PDF_GENERATED",
                  user["id"],
                  {"invoice_number": invoice_num,
                   "total_ttc":      total_ttc,
                   "hash":           doc_hash})
        return pdf_bytes

    except ImportError as e:
        raise RuntimeError(
            f"reportlab required for invoice PDF generation: {e}"
        )