# core/financial_engine.py
# ── PLAGENOR 4.0 — Financial Engine ──────────────────────────────────────────
# Handles all financial operations:
#   - Quote calculation (IBTIKAR budget + GENOCLAB billing)
#   - Invoice generation (number, amounts, VAT, document)
#   - Invoice PDF/DOCX generation via python-docx
#   - Payment recording
#   - Invoice integrity hashing
#   - Budget cap enforcement (IBTIKAR)
#   - Revenue summary, VAT summary, overdue tracking
#
# No Streamlit calls — pure business logic only.
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import hashlib
import uuid
import subprocess
from datetime import datetime, timedelta
from typing import Optional, Any

import config
from core.exceptions import (
    FinancialError,
    InvoiceAlreadyExistsError,
    InvoiceNotFoundError,
    InvalidQuoteAmountError,
    BudgetExceededError,
    DocumentGenerationError,
    RequestNotFoundError,
)
from core.repository import (
    get_invoice,
    get_invoice_by_request_id,
    get_all_invoices,
    save_invoice,
    next_invoice_number,
    get_request,
    save_request,
    get_service,
    create_notification,
    create_document_record,
)


# ── Constants ─────────────────────────────────────────────────────────────────
VAT_RATE           = float(getattr(config, "VAT_RATE",            0.19))
IBTIKAR_BUDGET_CAP = float(getattr(config, "IBTIKAR_BUDGET_CAP",  200_000.0))
CHANNEL_IBTIKAR    = config.CHANNEL_IBTIKAR
CHANNEL_GENOCLAB   = config.CHANNEL_GENOCLAB
INVOICES_DIR       = getattr(config, "INVOICES_DIR",  "data/invoices_pdf")
REPORTS_DIR        = getattr(config, "REPORTS_DIR",   "data/reports")
PLATFORM_NAME      = getattr(config, "PLATFORM_NAME", "PLAGENOR 4.0")
PLATFORM_INSTITUTION = getattr(
    config, "PLATFORM_INSTITUTION",
    "ESSBO — École Supérieure des Sciences Biologiques d'Oran"
)
PLATFORM_ADDRESS   = getattr(config, "PLATFORM_ADDRESS", "Oran, Algérie")
PLATFORM_EMAIL     = getattr(config, "PLATFORM_EMAIL",   "contact@essbo.dz")
PLATFORM_PHONE     = getattr(config, "PLATFORM_PHONE",   "")
INVOICE_PREFIX     = getattr(config, "INVOICE_PREFIX",   "GENOCLAB-INV")


# ── Utility helpers ───────────────────────────────────────────────────────────
def _now_iso() -> str:
    return datetime.utcnow().isoformat()


def _fmt_currency(amount: float) -> str:
    return f"{amount:,.2f} DZD"


def _round2(val: float) -> float:
    return round(val, 2)


def _fmt_date_fr(iso: str) -> str:
    """Returns DD/MM/YYYY from ISO timestamp."""
    if not iso:
        return "–"
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00"))
        return dt.strftime("%d/%m/%Y")
    except Exception:
        return iso[:10]


def _ensure_dir(path: str) -> None:
    os.makedirs(path, exist_ok=True)


def _compute_invoice_hash(invoice: dict) -> str:
    """
    Computes a SHA-256 integrity hash over immutable financial fields.
    Fields: invoice_number, request_id, total_ht, total_vat, total_ttc, created_at.
    """
    payload = (
        f"{invoice.get('invoice_number', '')}|"
        f"{invoice.get('request_id',     '')}|"
        f"{invoice.get('total_ht',        0):.2f}|"
        f"{invoice.get('total_vat',       0):.2f}|"
        f"{invoice.get('total_ttc',       0):.2f}|"
        f"{invoice.get('created_at',     '')}"
    )
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def verify_invoice_integrity(invoice: dict) -> bool:
    """Returns True if stored hash matches recomputed hash."""
    stored   = invoice.get("integrity_hash", "")
    computed = _compute_invoice_hash(invoice)
    return stored == computed


def _invoice_year(inv: dict) -> Optional[int]:
    try:
        return datetime.fromisoformat(
            inv.get("created_at", "").replace("Z", "+00:00")
        ).year
    except Exception:
        return None


def _invoice_month(inv: dict) -> Optional[int]:
    try:
        return datetime.fromisoformat(
            inv.get("created_at", "").replace("Z", "+00:00")
        ).month
    except Exception:
        return None


def _days_since(iso: str) -> int:
    if not iso:
        return 0
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00")).replace(tzinfo=None)
        return max((datetime.utcnow() - dt).days, 0)
    except Exception:
        return 0


# ══════════════════════════════════════════════════════════════════════════════
# QUOTE CALCULATION
# ══════════════════════════════════════════════════════════════════════════════

def calculate_quote(
    service_id:   str,
    form_data:    dict,
    channel:      str,
    override_ht:  Optional[float] = None,
) -> dict:
    """
    Calculates a quote for a request.

    Args:
        service_id:  Service being requested.
        form_data:   Request form_data dict (contains pricing inputs).
        channel:     CHANNEL_IBTIKAR or CHANNEL_GENOCLAB.
        override_ht: If provided, skips formula and uses this HT directly.

    Returns:
        {
            ht, vat, ttc, vat_rate,
            breakdown: [{ label, quantity, unit_price, subtotal }],
            notes: str,
        }
    """
    if override_ht is not None:
        if override_ht < 0:
            raise InvalidQuoteAmountError(override_ht)
        ht  = _round2(override_ht)
        vat = _round2(ht * VAT_RATE) if channel == CHANNEL_GENOCLAB else 0.0
        ttc = _round2(ht + vat)
        return {
            "ht":        ht,
            "vat":       vat,
            "ttc":       ttc,
            "vat_rate":  VAT_RATE if channel == CHANNEL_GENOCLAB else 0.0,
            "breakdown": [{"label": "Forfait", "quantity": 1,
                           "unit_price": ht, "subtotal": ht}],
            "notes":     "Montant saisi manuellement.",
        }

    service    = get_service(service_id)
    base_price = float((service or {}).get("base_price", 0))
    pricing    = form_data.get("pricing", {})

    # ── GENOCLAB — sample-based pricing ──────────────────────────────────────
    if channel == CHANNEL_GENOCLAB:
        samples   = max(int(pricing.get("samples",  1)), 1)
        runs      = max(int(pricing.get("runs",     1)), 1)
        extras    = float(pricing.get("extras",     0))
        urgency   = float(pricing.get("urgency_factor", 1.0))

        sample_sub = _round2(base_price * samples)
        run_sub    = _round2(base_price * 0.10 * runs)
        extras_sub = _round2(extras)
        subtotal   = _round2(
            (sample_sub + run_sub + extras_sub) * urgency
        )

        breakdown = []
        if sample_sub:
            breakdown.append({
                "label":      "Analyse par échantillon",
                "quantity":   samples,
                "unit_price": base_price,
                "subtotal":   sample_sub,
            })
        if run_sub:
            breakdown.append({
                "label":      "Frais de run",
                "quantity":   runs,
                "unit_price": _round2(base_price * 0.10),
                "subtotal":   run_sub,
            })
        if extras_sub:
            breakdown.append({
                "label":      "Extras / consommables",
                "quantity":   1,
                "unit_price": extras_sub,
                "subtotal":   extras_sub,
            })
        if urgency != 1.0:
            breakdown.append({
                "label":      f"Facteur urgence ×{urgency:.2f}",
                "quantity":   1,
                "unit_price": 0,
                "subtotal":   _round2(subtotal - subtotal / urgency),
            })

        ht  = subtotal
        vat = _round2(ht * VAT_RATE)
        ttc = _round2(ht + vat)
        return {
            "ht":        ht,
            "vat":       vat,
            "ttc":       ttc,
            "vat_rate":  VAT_RATE,
            "breakdown": breakdown,
            "notes":     f"Calcul automatique — {samples} échantillon(s).",
        }

    # ── IBTIKAR — budget-based (no VAT) ──────────────────────────────────────
    else:
        budget    = form_data.get("budget", {})
        requested = float(budget.get("requested", base_price))

        if requested > IBTIKAR_BUDGET_CAP:
            raise BudgetExceededError(requested, IBTIKAR_BUDGET_CAP)

        return {
            "ht":        _round2(requested),
            "vat":       0.0,
            "ttc":       _round2(requested),
            "vat_rate":  0.0,
            "breakdown": [{
                "label":      "Budget de recherche demandé",
                "quantity":   1,
                "unit_price": requested,
                "subtotal":   requested,
            }],
            "notes": "Financement IBTIKAR-DGRSDT — TVA non applicable.",
        }


def set_request_quote(
    request_id:  str,
    amount_ht:   float,
    actor:       dict,
    notes:       str = "",
    quote_notes: str = "",
) -> dict:
    """
    Sets (or updates) the quote_amount on a request.
    Does NOT change workflow state — call workflow_engine.transition()
    to advance to QUOTE_SENT after calling this.

    Returns the updated request dict.
    """
    if amount_ht < 0:
        raise InvalidQuoteAmountError(amount_ht)

    req = get_request(request_id)
    if not req:
        raise RequestNotFoundError(request_id)

    vat_rate = VAT_RATE if req.get("channel") == CHANNEL_GENOCLAB else 0.0
    vat      = _round2(amount_ht * vat_rate)
    ttc      = _round2(amount_ht + vat)

    req["quote_amount"]   = _round2(amount_ht)
    req["quote_vat"]      = vat
    req["quote_ttc"]      = ttc
    req["quote_vat_rate"] = vat_rate
    req["quote_notes"]    = quote_notes
    req["quote_set_by"]   = actor.get("id",  "system")
    req["quote_set_at"]   = _now_iso()
    req["updated_at"]     = _now_iso()

    if notes:
        req.setdefault("notes", [])
        if isinstance(req["notes"], list):
            req["notes"].append({
                "text":      notes,
                "timestamp": _now_iso(),
                "author_id": actor.get("id", "system"),
            })

    save_request(req)
    return req


# ══════════════════════════════════════════════════════════════════════════════
# INVOICE GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def generate_invoice(
    request: dict,
    actor:   dict,
    force:   bool = False,
) -> dict:
    """
    Generates a complete invoice record for a request.

    Args:
        request: Full request dict (must have quote_amount).
        actor:   User performing the action.
        force:   If True, regenerates even if invoice already exists.

    Returns:
        The saved invoice dict.

    Raises:
        InvoiceAlreadyExistsError — invoice exists and force=False.
        InvalidQuoteAmountError   — quote_amount is 0 or missing.
    """
    request_id = request.get("id", "")

    # ── Guard: already exists ─────────────────────────────────────────────────
    existing = get_invoice_by_request_id(request_id)
    if existing and not force:
        raise InvoiceAlreadyExistsError(request_id)

    # ── Guard: quote required ─────────────────────────────────────────────────
    quote_ht = float(request.get("quote_amount", 0))
    if quote_ht <= 0:
        raise InvalidQuoteAmountError(quote_ht)

    channel  = request.get("channel", CHANNEL_GENOCLAB)
    vat_rate = VAT_RATE if channel == CHANNEL_GENOCLAB else 0.0
    vat      = _round2(quote_ht * vat_rate)
    ttc      = _round2(quote_ht + vat)

    form_data = request.get("form_data", {})
    requester = form_data.get("requester", {})
    pricing   = form_data.get("pricing",   {})

    # ── Build line items ──────────────────────────────────────────────────────
    line_items = request.get("quote_breakdown", [])
    if not line_items:
        service  = get_service(request.get("service_id", ""))
        svc_name = (service or {}).get("name", "Service d'analyse")
        samples  = max(int(pricing.get("samples", 1)), 1)
        line_items = [{
            "label":      svc_name,
            "quantity":   samples,
            "unit_price": _round2(quote_ht / samples),
            "subtotal":   quote_ht,
            "vat_rate":   vat_rate,
        }]

    # ── Invoice number ────────────────────────────────────────────────────────
    inv_number = (
        existing.get("invoice_number")
        if (existing and force)
        else next_invoice_number()
    )

    # ── Build invoice dict ────────────────────────────────────────────────────
    invoice = {
        "id":              existing["id"] if (existing and force) else str(uuid.uuid4()),
        "invoice_number":  inv_number,
        "request_id":      request_id,
        "channel":         channel,
        "service_id":      request.get("service_id", ""),

        # Amounts
        "total_ht":        quote_ht,
        "total_vat":       vat,
        "total_ttc":       ttc,
        "vat_rate":        vat_rate,
        "line_items":      line_items,

        # Client snapshot
        "client_name":     requester.get("full_name",   request.get("client_name",   "–")),
        "client_email":    requester.get("email",       request.get("client_email",   "–")),
        "client_phone":    requester.get("phone",       "–"),
        "institution":     requester.get("institution", request.get("institution",    "–")),
        "organization_id": request.get("organization_id", ""),

        # Payment state
        "paid":            False,
        "paid_at":         None,
        "payment_ref":     None,
        "payment_note":    None,

        # Metadata
        "created_at":      _now_iso(),
        "created_by":      actor.get("id", "system"),
        "updated_at":      _now_iso(),
        "pdf_path":        None,
        "integrity_hash":  None,
        "hash_short":      None,
        "generation_warnings": [],
    }

    # ── Integrity hash ────────────────────────────────────────────────────────
    h                         = _compute_invoice_hash(invoice)
    invoice["integrity_hash"] = h
    invoice["hash_short"]     = h[:12]

    # ── Save invoice ──────────────────────────────────────────────────────────
    save_invoice(invoice)

    # ── Generate document (non-blocking) ─────────────────────────────────────
    try:
        pdf_path = generate_invoice_pdf(invoice, request)
        if pdf_path:
            invoice["pdf_path"] = pdf_path
            save_invoice(invoice)
            create_document_record(
                request_id = request_id,
                filename   = os.path.basename(pdf_path),
                path       = pdf_path,
                doc_type   = "INVOICE",
                created_by = actor.get("id", "system"),
            )
    except Exception as e:
        invoice["generation_warnings"].append(str(e))
        save_invoice(invoice)

    # ── Stamp invoice_id on request ───────────────────────────────────────────
    request["invoice_id"]     = invoice["id"]
    request["invoice_number"] = inv_number
    request["updated_at"]     = _now_iso()
    save_request(request)

    return invoice


# ══════════════════════════════════════════════════════════════════════════════
# PDF / DOCX GENERATION
# ══════════════════════════════════════════════════════════════════════════════

def generate_invoice_pdf(
    invoice: dict,
    request: dict,
) -> Optional[str]:
    """
    Generates an invoice document using python-docx.
    Attempts LibreOffice conversion to PDF if available.
    Falls back to .docx if LibreOffice is not installed.

    Returns:
        Absolute path to the generated file (.pdf or .docx), or None on error.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        raise DocumentGenerationError(
            "python-docx non installé. "
            "Exécutez: pip install python-docx"
        )

    _ensure_dir(INVOICES_DIR)

    inv_number = invoice.get("invoice_number", invoice["id"][:8])
    safe_num   = inv_number.replace("/", "-").replace("\\", "-")
    docx_path  = os.path.join(INVOICES_DIR, f"facture_{safe_num}.docx")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def _add_heading(text: str, level: int = 1) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold      = True
        run.font.size = Pt(16 - (level - 1) * 2)
        run.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    def _add_line(label: str, value: str, bold_val: bool = False) -> None:
        p   = doc.add_paragraph()
        r1  = p.add_run(f"{label}: ")
        r1.bold = True
        r2  = p.add_run(value)
        r2.bold = bold_val

    # ── Header ────────────────────────────────────────────────────────────────
    _add_heading(PLATFORM_NAME)
    _add_heading(PLATFORM_INSTITUTION, level=2)
    doc.add_paragraph(PLATFORM_ADDRESS).alignment = WD_ALIGN_PARAGRAPH.CENTER
    if PLATFORM_EMAIL:
        doc.add_paragraph(
            f"✉ {PLATFORM_EMAIL}"
            + (f"  |  ☎ {PLATFORM_PHONE}" if PLATFORM_PHONE else "")
        ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("─" * 80)

    # ── Invoice title ─────────────────────────────────────────────────────────
    title_p           = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_r           = title_p.add_run(f"FACTURE  N°  {inv_number}")
    title_r.bold      = True
    title_r.font.size = Pt(14)
    title_r.font.color.rgb = RGBColor(0x1A, 0xBC, 0x9C)

    doc.add_paragraph(
        f"Date d'émission: {_fmt_date_fr(invoice.get('created_at', ''))}"
    )
    doc.add_paragraph("─" * 80)

    # ── Client section ────────────────────────────────────────────────────────
    doc.add_heading("FACTURÉ À", level=2)
    _add_line("Client",      invoice.get("client_name",  "–"))
    _add_line("Institution", invoice.get("institution",  "–"))
    _add_line("Email",       invoice.get("client_email", "–"))
    if invoice.get("client_phone", "–") != "–":
        _add_line("Téléphone", invoice.get("client_phone", "–"))

    doc.add_paragraph("")

    # ── Request info ──────────────────────────────────────────────────────────
    doc.add_heading("DÉTAILS DE LA DEMANDE", level=2)
    _add_line("Référence demande", request.get("id", "–")[:8].upper())
    _add_line("Canal",            invoice.get("channel", "–"))

    service = get_service(request.get("service_id", ""))
    if service:
        _add_line("Service", service.get("name", "–"))

    doc.add_paragraph("")

    # ── Line items table ──────────────────────────────────────────────────────
    doc.add_heading("DÉTAIL DES PRESTATIONS", level=2)

    line_items = invoice.get("line_items", [])
    if line_items:
        tbl = doc.add_table(
            rows  = 1 + len(line_items),
            cols  = 4,
        )
        tbl.style = "Table Grid"

        # Header row
        hdr_cells = tbl.rows[0].cells
        for i, h in enumerate(["Désignation", "Qté", "Prix unitaire (DZD)", "Sous-total (DZD)"]):
            hdr_cells[i].text = h
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Background colour via XML (blue header)
            from docx.oxml.ns import qn
            from docx.oxml    import OxmlElement
            tc_pr  = hdr_cells[i]._tc.get_or_add_tcPr()
            shd    = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "1B4F72")
            tc_pr.append(shd)

        # Data rows
        for j, item in enumerate(line_items, start=1):
            cells = tbl.rows[j].cells
            cells[0].text = str(item.get("label",      "–"))
            cells[1].text = str(item.get("quantity",    1))
            cells[2].text = f"{float(item.get('unit_price', 0)):,.2f}"
            cells[3].text = f"{float(item.get('subtotal',   0)):,.2f}"

    doc.add_paragraph("")

    # ── Financial summary ─────────────────────────────────────────────────────
    doc.add_heading("RÉCAPITULATIF FINANCIER", level=2)

    summary_tbl = doc.add_table(rows=4, cols=2)
    summary_tbl.style = "Table Grid"

    vat_rate_pct = float(invoice.get("vat_rate", 0)) * 100
    rows_data = [
        ("Total HT",                   f"{float(invoice.get('total_ht',  0)):,.2f} DZD"),
        (f"TVA ({vat_rate_pct:.0f}%)", f"{float(invoice.get('total_vat', 0)):,.2f} DZD"),
        ("Total TTC",                  f"{float(invoice.get('total_ttc', 0)):,.2f} DZD"),
        ("Statut paiement",            "✅ Réglée" if invoice.get("paid") else "⏳ En attente"),
    ]

    for i, (label, value) in enumerate(rows_data):
        r = summary_tbl.rows[i]
        r.cells[0].text = label
        r.cells[1].text = value
        for para in r.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
        # Highlight TTC row
        if label == "Total TTC":
            for cell in r.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0x1A, 0xBC, 0x9C)

    doc.add_paragraph("")

    # ── Payment info if already paid ──────────────────────────────────────────
    if invoice.get("paid"):
        doc.add_heading("INFORMATION DE PAIEMENT", level=2)
        _add_line("Date de paiement",  _fmt_date_fr(invoice.get("paid_at", "")))
        _add_line("Référence",         invoice.get("payment_ref", "–"))
        if invoice.get("payment_note"):
            _add_line("Note",          invoice.get("payment_note", ""))

    # ── Integrity footer ──────────────────────────────────────────────────────
    doc.add_paragraph("─" * 80)
    footer_p = doc.add_paragraph(
        f"Hash d'intégrité: {invoice.get('hash_short', '–')}  |  "
        f"Émis le: {_fmt_date_fr(invoice.get('created_at', ''))}  |  "
        f"{PLATFORM_NAME} v{getattr(config, 'PLATFORM_VERSION', '4.0')}"
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # ── Notes ─────────────────────────────────────────────────────────────────
    if invoice.get("quote_notes") or (
        invoice.get("channel") == CHANNEL_IBTIKAR
    ):
        doc.add_paragraph("")
        doc.add_heading("NOTE", level=2)
        note_text = (
            invoice.get("quote_notes")
            or "Financement institutionnel IBTIKAR-DGRSDT — TVA non applicable."
        )
        doc.add_paragraph(note_text)

    # ── Save .docx ────────────────────────────────────────────────────────────
    doc.save(docx_path)

    # ── Attempt LibreOffice conversion ────────────────────────────────────────
    pdf_path = _convert_docx_to_pdf(docx_path)
    return pdf_path if pdf_path else docx_path


def _convert_docx_to_pdf(docx_path: str) -> Optional[str]:
    """
    Tries to convert a .docx to .pdf using LibreOffice headless.
    Returns the PDF path if successful, None if LibreOffice is not available.
    """
    lo_commands = [
        "libreoffice",
        "libreoffice7.6",
        "soffice",
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]

    out_dir = os.path.dirname(docx_path)

    for cmd in lo_commands:
        try:
            result = subprocess.run(
                [
                    cmd,
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", out_dir,
                    docx_path,
                ],
                capture_output = True,
                timeout        = 30,
            )
            if result.returncode == 0:
                pdf_path = docx_path.replace(".docx", ".pdf")
                if os.path.exists(pdf_path):
                    return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue
        except Exception:
            continue

    return None  # LibreOffice not available — caller falls back to .docx


def regenerate_invoice_pdf(
    invoice_id: str,
    actor:      dict,
) -> Optional[str]:
    """
    Regenerates the PDF/DOCX for an existing invoice.
    Updates invoice.pdf_path and document record.

    Returns:
        Path to the regenerated file, or None.

    Raises:
        InvoiceNotFoundError — invoice_id not found.
    """
    invoice = get_invoice(invoice_id)
    if not invoice:
        raise InvoiceNotFoundError(invoice_id)

    request_id = invoice.get("request_id", "")
    request    = get_request(request_id) if request_id else {}

    try:
        path = generate_invoice_pdf(invoice, request or {})
    except Exception as e:
        raise DocumentGenerationError(str(e))

    if path:
        invoice["pdf_path"]   = path
        invoice["updated_at"] = _now_iso()
        save_invoice(invoice)

        if request_id:
            try:
                create_document_record(
                    request_id = request_id,
                    filename   = os.path.basename(path),
                    path       = path,
                    doc_type   = "INVOICE",
                    created_by = actor.get("id", "system"),
                )
            except Exception:
                pass

        try:
            from core.audit_engine import log_invoice_event
            log_invoice_event(
                invoice_id     = invoice_id,
                event          = "PDF_GENERATED",
                actor          = actor,
                invoice_number = invoice.get("invoice_number", ""),
                amount_ttc     = float(invoice.get("total_ttc", 0)),
                details        = f"Fichier: {os.path.basename(path)}",
            )
        except Exception:
            pass

    return path


# ══════════════════════════════════════════════════════════════════════════════
# PAYMENT RECORDING
# ══════════════════════════════════════════════════════════════════════════════

def record_payment(
    invoice_id:  str,
    paid_date:   str,
    payment_ref: str,
    actor:       dict,
    notes:       str = "",
) -> dict:
    """
    Records a payment against an invoice.

    Args:
        invoice_id:  Invoice to mark as paid.
        paid_date:   Payment date string (YYYY-MM-DD or ISO).
        payment_ref: Bank transfer / cheque reference.
        actor:       User recording the payment.
        notes:       Optional payment note.

    Returns:
        Updated invoice dict.

    Raises:
        InvoiceNotFoundError — invoice not found.
        FinancialError       — already paid.
    """
    invoice = get_invoice(invoice_id)
    if not invoice:
        raise InvoiceNotFoundError(invoice_id)

    if invoice.get("paid"):
        raise FinancialError(
            f"La facture {invoice.get('invoice_number', invoice_id[:8])} "
            f"est déjà marquée comme réglée "
            f"(le {_fmt_date_fr(invoice.get('paid_at', ''))}).",
            "ALREADY_PAID",
        )

    if not payment_ref or not payment_ref.strip():
        raise FinancialError(
            "La référence de paiement est obligatoire.",
            "MISSING_PAYMENT_REF",
        )

    # ── Update invoice ────────────────────────────────────────────────────────
    invoice["paid"]         = True
    invoice["paid_at"]      = (
        paid_date
        if "T" in paid_date
        else f"{paid_date}T00:00:00"
    )
    invoice["payment_ref"]  = payment_ref.strip()
    invoice["payment_note"] = notes.strip() if notes else None
    invoice["paid_by"]      = actor.get("id", "system")
    invoice["updated_at"]   = _now_iso()

    save_invoice(invoice)

    # ── Audit log ─────────────────────────────────────────────────────────────
    try:
        from core.audit_engine import log_invoice_event
        log_invoice_event(
            invoice_id     = invoice_id,
            event          = "INVOICE_PAID",
            actor          = actor,
            invoice_number = invoice.get("invoice_number", ""),
            amount_ttc     = float(invoice.get("total_ttc", 0)),
            details        = (
                f"Réf: {payment_ref} | "
                f"Date: {_fmt_date_fr(invoice['paid_at'])}"
            ),
        )
    except Exception:
        pass

    # ── Notify finance role ───────────────────────────────────────────────────
    try:
        create_notification(
            title   = "💳 Paiement enregistré",
            message = (
                f"Facture {invoice.get('invoice_number','–')} "
                f"marquée comme réglée — "
                f"{_fmt_currency(float(invoice.get('total_ttc', 0)))}."
            ),
            level   = "success",
            role    = config.ROLE_FINANCE,
        )
    except Exception:
        pass

    return invoice


# ══════════════════════════════════════════════════════════════════════════════
# REVENUE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

def get_revenue_summary(year: Optional[int] = None) -> dict:
    """
    Returns a complete revenue summary for a given year (or all time).

    Returns:
        {
            year,
            total_invoices,
            paid_invoices,
            unpaid_invoices,
            total_ht,
            total_vat,
            total_ttc,
            encaisse_ttc,
            a_encaisser_ttc,
            collection_rate_pct,
            by_month:   { month_int: { ht, vat, ttc, paid_ttc, count } },
            by_channel: { channel: { count, total_ttc, paid_ttc } },
        }
    """
    all_invs = get_all_invoices()

    if year:
        invs = [i for i in all_invs if _invoice_year(i) == year]
    else:
        invs = list(all_invs)

    paid_invs   = [i for i in invs if i.get("paid")]
    unpaid_invs = [i for i in invs if not i.get("paid")]

    total_ht    = _round2(sum(float(i.get("total_ht",  0)) for i in invs))
    total_vat   = _round2(sum(float(i.get("total_vat", 0)) for i in invs))
    total_ttc   = _round2(sum(float(i.get("total_ttc", 0)) for i in invs))
    enc_ttc     = _round2(sum(float(i.get("total_ttc", 0)) for i in paid_invs))
    a_enc_ttc   = _round2(sum(float(i.get("total_ttc", 0)) for i in unpaid_invs))
    coll_rate   = round(enc_ttc / total_ttc * 100, 1) if total_ttc else 0.0

    # ── Monthly breakdown ─────────────────────────────────────────────────────
    by_month: dict = {}
    for m in range(1, 13):
        m_invs = [i for i in invs if _invoice_month(i) == m]
        by_month[m] = {
            "ht":       _round2(sum(float(i.get("total_ht",  0)) for i in m_invs)),
            "vat":      _round2(sum(float(i.get("total_vat", 0)) for i in m_invs)),
            "ttc":      _round2(sum(float(i.get("total_ttc", 0)) for i in m_invs)),
            "paid_ttc": _round2(sum(
                float(i.get("total_ttc", 0))
                for i in m_invs if i.get("paid")
            )),
            "count":    len(m_invs),
        }

    # ── Channel breakdown ─────────────────────────────────────────────────────
    by_channel: dict = {}
    for ch in [CHANNEL_IBTIKAR, CHANNEL_GENOCLAB]:
        ch_invs = [i for i in invs if i.get("channel") == ch]
        by_channel[ch] = {
            "count":     len(ch_invs),
            "total_ht":  _round2(sum(float(i.get("total_ht",  0)) for i in ch_invs)),
            "total_ttc": _round2(sum(float(i.get("total_ttc", 0)) for i in ch_invs)),
            "paid_ttc":  _round2(sum(
                float(i.get("total_ttc", 0))
                for i in ch_invs if i.get("paid")
            )),
        }

    return {
        "year":                year,
        "total_invoices":      len(invs),
        "paid_invoices":       len(paid_invs),
        "unpaid_invoices":     len(unpaid_invs),
        "total_ht":            total_ht,
        "total_vat":           total_vat,
        "total_ttc":           total_ttc,
        "encaisse_ttc":        enc_ttc,
        "a_encaisser_ttc":     a_enc_ttc,
        "collection_rate_pct": coll_rate,
        "by_month":            by_month,
        "by_channel":          by_channel,
    }


# ══════════════════════════════════════════════════════════════════════════════
# VAT SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

def get_vat_summary(year: Optional[int] = None) -> dict:
    """
    Returns VAT summary for GENOCLAB invoices only (IBTIKAR = VAT exempt).

    Returns:
        {
            year,
            annual: { count, ht, vat, ttc },
            quarters: {
                "T1": { count, ht, vat, ttc },
                "T2": { count, ht, vat, ttc },
                "T3": { count, ht, vat, ttc },
                "T4": { count, ht, vat, ttc },
            },
        }
    """
    all_invs = get_all_invoices()

    # VAT applies only to GENOCLAB
    invs = [
        i for i in all_invs
        if i.get("channel") == CHANNEL_GENOCLAB
        and (year is None or _invoice_year(i) == year)
    ]

    def _quarter(m: Optional[int]) -> str:
        if not m:
            return "T4"
        return {1: "T1", 2: "T1", 3: "T1",
                4: "T2", 5: "T2", 6: "T2",
                7: "T3", 8: "T3", 9: "T3",
                10: "T4", 11: "T4", 12: "T4"}.get(m, "T4")

    quarters: dict = {
        "T1": {"count": 0, "ht": 0.0, "vat": 0.0, "ttc": 0.0},
        "T2": {"count": 0, "ht": 0.0, "vat": 0.0, "ttc": 0.0},
        "T3": {"count": 0, "ht": 0.0, "vat": 0.0, "ttc": 0.0},
        "T4": {"count": 0, "ht": 0.0, "vat": 0.0, "ttc": 0.0},
    }

    for inv in invs:
        q = _quarter(_invoice_month(inv))
        quarters[q]["count"] += 1
        quarters[q]["ht"]    += float(inv.get("total_ht",  0))
        quarters[q]["vat"]   += float(inv.get("total_vat", 0))
        quarters[q]["ttc"]   += float(inv.get("total_ttc", 0))

    # Round quarterly totals
    for q in quarters:
        for k in ["ht", "vat", "ttc"]:
            quarters[q][k] = _round2(quarters[q][k])

    annual = {
        "count": len(invs),
        "ht":    _round2(sum(float(i.get("total_ht",  0)) for i in invs)),
        "vat":   _round2(sum(float(i.get("total_vat", 0)) for i in invs)),
        "ttc":   _round2(sum(float(i.get("total_ttc", 0)) for i in invs)),
    }

    return {
        "year":     year,
        "annual":   annual,
        "quarters": quarters,
    }


# ══════════════════════════════════════════════════════════════════════════════
# OVERDUE INVOICES
# ══════════════════════════════════════════════════════════════════════════════

def get_overdue_invoices(overdue_days: int = 30) -> list[dict]:
    """
    Returns unpaid invoices older than overdue_days.
    Enriches each invoice with days_overdue field.

    Args:
        overdue_days: Age threshold in days (default 30).

    Returns:
        List of enriched invoice dicts, sorted by days_overdue DESC.
    """
    all_invs = get_all_invoices()
    result   = []

    for inv in all_invs:
        if inv.get("paid"):
            continue
        days = _days_since(inv.get("created_at", ""))
        if days >= overdue_days:
            enriched               = dict(inv)
            enriched["days_overdue"] = days
            enriched["overdue"]      = True
            result.append(enriched)

    return sorted(result, key=lambda i: i["days_overdue"], reverse=True)


# ══════════════════════════════════════════════════════════════════════════════
# PLATFORM NOTE (IBTIKAR — auto-generated on VALIDATED transition)
# ══════════════════════════════════════════════════════════════════════════════

def generate_platform_note(
    request: dict,
    actor:   dict,
) -> Optional[str]:
    """
    Generates an institutional validation note (NOTE-PLATEFORME) for
    IBTIKAR requests when they reach VALIDATED state.

    Returns:
        Path to the generated .docx file, or None on failure.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    _ensure_dir(REPORTS_DIR)

    prefix   = getattr(config, "PLATFORM_NOTE_PREFIX", "NOTE-PLATEFORME")
    req_id   = request.get("id", str(uuid.uuid4()))[:8].upper()
    filename = f"{prefix}-{req_id}-{datetime.utcnow().strftime('%Y%m%d')}.docx"
    path     = os.path.join(REPORTS_DIR, filename)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PLATFORM_NAME)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_paragraph(
        PLATFORM_INSTITUTION
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("─" * 60)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run(f"NOTE DE VALIDATION — N° {prefix}-{req_id}")
    tr.bold = True
    tr.font.size = Pt(13)

    doc.add_paragraph(
        f"Date: {_fmt_date_fr(_now_iso())}"
    )
    doc.add_paragraph("─" * 60)

    # Request details
    doc.add_heading("DEMANDE", level=2)
    form_data = request.get("form_data", {})
    requester = form_data.get("requester", {})
    budget    = form_data.get("budget",    {})

    for label, value in [
        ("Référence",      req_id),
        ("Canal",          CHANNEL_IBTIKAR),
        ("Demandeur",      requester.get("full_name",   "–")),
        ("Institution",    requester.get("institution", "–")),
        ("Email",          requester.get("email",       "–")),
        ("Budget demandé", f"{float(budget.get('requested', 0)):,.0f} DZD"),
    ]:
        p2  = doc.add_paragraph()
        p2.add_run(f"{label}: ").bold = True
        p2.add_run(str(value))

    doc.add_paragraph("")
    doc.add_heading("DÉCISION", level=2)
    dec_p = doc.add_paragraph()
    dec_r = dec_p.add_run("✅ DEMANDE VALIDÉE")
    dec_r.bold = True
    dec_r.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

    doc.add_paragraph(
        "La présente demande a été examinée et validée par l'administration "
        "de la plateforme PLAGENOR. Elle est transmise pour approbation "
        "en vue d'un financement IBTIKAR-DGRSDT."
    )

    doc.add_paragraph("─" * 60)
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_p.add_run(
        f"Validé par: {actor.get('username', 'Admin')}  |  "
        f"{_fmt_date_fr(_now_iso())}"
    ).font.size = Pt(9)

    doc.save(path)

    # Register document
    try:
        create_document_record(
            request_id = request.get("id", ""),
            filename   = filename,
            path       = path,
            doc_type   = "PLATFORM_NOTE",
            created_by = actor.get("id", "system"),
        )
    except Exception:
        pass

    return path


# ══════════════════════════════════════════════════════════════════════════════
# REPORT DOCX TEMPLATE
# ══════════════════════════════════════════════════════════════════════════════

def generate_report_docx(
    request: dict,
    actor:   dict,
) -> Optional[str]:
    """
    Generates an empty analysis report template (.docx) for a request.
    Called automatically by workflow_engine on REPORT_UPLOADED hook.
    Members replace the content with actual results.

    Returns:
        Path to the generated .docx, or None on failure.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return None

    _ensure_dir(REPORTS_DIR)

    req_id   = request.get("id", str(uuid.uuid4()))[:8].upper()
    channel  = request.get("channel", "")
    filename = f"rapport_{channel}_{req_id}_{datetime.utcnow().strftime('%Y%m%d')}.docx"
    path     = os.path.join(REPORTS_DIR, filename)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # Header
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = h.add_run(f"RAPPORT D'ANALYSE — {PLATFORM_NAME}")
    hr.bold = True
    hr.font.size = Pt(15)
    hr.font.color.rgb = RGBColor(0x1B, 0x4F, 0x72)

    doc.add_paragraph(
        PLATFORM_INSTITUTION
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("─" * 60)

    # Request info
    form_data = request.get("form_data", {})
    requester = form_data.get("requester", {})
    service   = get_service(request.get("service_id", ""))

    for label, value in [
        ("Référence demande", req_id),
        ("Canal",             channel),
        ("Service",           (service or {}).get("name", "–")),
        ("Demandeur / Client",requester.get("full_name", "–")),
        ("Institution",       requester.get("institution", "–")),
        ("Date du rapport",   _fmt_date_fr(_now_iso())),
        ("Analyste assigné",  request.get("assigned_member_name", "–")),
    ]:
        p2 = doc.add_paragraph()
        p2.add_run(f"{label}: ").bold = True
        p2.add_run(str(value))

    doc.add_paragraph("")

    # Sections
    for section_title in [
        "1. Résumé exécutif",
        "2. Matériels et méthodes",
        "3. Résultats",
        "4. Discussion",
        "5. Conclusion",
        "6. Annexes",
    ]:
        doc.add_heading(section_title, level=2)
        doc.add_paragraph("[À compléter par l'analyste]")
        doc.add_paragraph("")

    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph(
        f"Document généré automatiquement par {PLATFORM_NAME} | "
        f"{_fmt_date_fr(_now_iso())}"
    )
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    doc.save(path)

    try:
        create_document_record(
            request_id = request.get("id", ""),
            filename   = filename,
            path       = path,
            doc_type   = "REPORT_TEMPLATE",
            created_by = actor.get("id", "system"),
        )
    except Exception:
        pass

    return path